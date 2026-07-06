from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 頁面設定：A4，上下左右邊距 1.5cm ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)

# ── 顏色定義 ──
DARK_GREEN  = RGBColor(0x1B, 0x5E, 0x20)   # 深綠（標題）
MID_GREEN   = RGBColor(0x2E, 0x7D, 0x32)   # 中綠（表頭）
LIGHT_GREEN = RGBColor(0xE8, 0xF5, 0xE9)   # 淺綠（表頭底色）
ACCENT      = RGBColor(0xFF, 0x8F, 0x00)   # 琥珀（強調數字）
GRAY_BG     = RGBColor(0xF5, 0xF5, 0xF5)   # 淺灰（交替行）
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TEXT        = RGBColor(0x21, 0x21, 0x21)

# ── Helper ──
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

def add_section_title(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GREEN
    # 底線用段落border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1B5E20')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_kv_line(doc, label, value, value_color=TEXT):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{label}：')
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2 = p.add_run(value)
    r2.bold = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = value_color

# ══════════════════════════════════════
#  標題區塊
# ══════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(2)
r = title_p.add_run('十全特好食品｜行銷月報')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARK_GREEN

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(4)
r2 = sub_p.add_run('2026 年 6 月  |  官網・KOL・Meta 廣告 整合分析')
r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

# ══════════════════════════════════════
#  一、官網銷售總覽
# ══════════════════════════════════════
add_section_title(doc, '一、官網銷售總覽（2026/06/01–06/30）')

# KPI 四格橫排
kpi_table = doc.add_table(rows=1, cols=4)
kpi_table.style = 'Table Grid'
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
kpis = [
    ('月淨營收', 'NT$ 204,034'),
    ('總訂單數', '188 筆'),
    ('平均客單價', 'NT$ 1,085'),
    ('單日最高', '06/09  NT$ 34,876'),
]
for i, (label, val) in enumerate(kpis):
    cell = kpi_table.rows[0].cells[i]
    cell.width = Cm(4.2)
    set_cell_bg(cell, LIGHT_GREEN)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = p.add_run(label + '\n')
    r_lbl.font.size = Pt(8); r_lbl.font.color.rgb = MID_GREEN
    r_val = p.add_run(val)
    r_val.bold = True; r_val.font.size = Pt(10); r_val.font.color.rgb = ACCENT

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# TOP 5 商品
add_kv_line(doc, '暢銷商品 TOP 5', '')
prod_table = doc.add_table(rows=6, cols=5)
prod_table.style = 'Table Grid'
headers = ['排名', '商品名稱', '銷售件數', '銷售金額', '佔營收%']
top5 = [
    ('1', 'YiBiYaYa 100% 綜合莓果汁', '47', 'NT$ 17,625', '8.2%'),
    ('2', 'YiBiYaYa 100億益生菌（30包）', '46', 'NT$ 16,634', '7.8%'),
    ('3', '濃縮醋禮盒（青梅＋蘋果）', '57', 'NT$ 14,820', '6.9%'),
    ('4', 'YiBiYaYa 100% 綜合蔬果汁', '36', 'NT$ 13,500', '6.3%'),
    ('5', 'YiBiYaYa 100% 蘋果胡蘿蔔汁', '31', 'NT$ 11,625', '5.4%'),
]
col_widths = [Cm(1.1), Cm(6.8), Cm(1.8), Cm(2.2), Cm(1.8)]
for j, h in enumerate(headers):
    cell = prod_table.rows[0].cells[j]
    cell.width = col_widths[j]
    set_cell_bg(cell, MID_GREEN)
    cell_text(cell, h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row_data in enumerate(top5):
    bg = LIGHT_GREEN if i % 2 == 0 else WHITE
    for j, val in enumerate(row_data):
        cell = prod_table.rows[i+1].cells[j]
        set_cell_bg(cell, bg)
        align = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
        cell_text(cell, val, size=8.5, align=align)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# 銷售洞察
p_insight = doc.add_paragraph()
p_insight.paragraph_format.space_before = Pt(2)
p_insight.paragraph_format.space_after  = Pt(3)
ri = p_insight.add_run('📌 洞察：')
ri.bold = True; ri.font.size = Pt(8.5); ri.font.color.rgb = DARK_GREEN
ri2 = p_insight.add_run(
    ' YiBiYaYa 系列（果汁+益生菌）共佔整體營收 35.8%，為主力產品線。'
    '06/09 單日訂單量達 33 筆、營收 NT$34,876，為全月最高峰，'
    '與 KOL 發文時間（06/08–06/09）高度吻合，顯示 KOL 引流成效顯著。'
)
ri2.font.size = Pt(8.5)

# ══════════════════════════════════════
#  二、KOL 合作成效
# ══════════════════════════════════════
add_section_title(doc, '二、KOL 合作成效（2026 年 6 月）')

add_kv_line(doc, '合作 KOL 數', '21 位', ACCENT)
add_kv_line(doc, '總合作費用', 'NT$ 181,500（不含費用未定者）', ACCENT)

kol_table = doc.add_table(rows=7, cols=4)
kol_table.style = 'Table Grid'
kol_headers = ['KOL 名稱', '粉絲規模', '費用 (NT$)', '觀看率']
kol_data = [
    ('萬事如翊',           '5,000–1萬',  '3,000',  '10.0%  ⭐'),
    ('Orly\'s 食研廚房',   '5萬–10萬',   '27,000', '6.0%   ⭐'),
    ('威媽的吃貨日常',     '5,000–1萬',  '3,000',  '2.0%'),
    ('長頸鹿學姊的健康自煮','5,000–1萬', '5,000',  '1.3%'),
    ('小廚娘阿貝',         '3,000–5,000','免費',   '1.1%'),
    ('100天便當計畫',      '10萬以上',   '75,000', '0.3%  ⚠'),
]
kol_col_w = [Cm(4.5), Cm(2.5), Cm(2.2), Cm(2.0)]
for j, h in enumerate(kol_headers):
    cell = kol_table.rows[0].cells[j]
    cell.width = kol_col_w[j]
    set_cell_bg(cell, MID_GREEN)
    cell_text(cell, h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row_data in enumerate(kol_data):
    bg = LIGHT_GREEN if i % 2 == 0 else WHITE
    for j, val in enumerate(row_data):
        cell = kol_table.rows[i+1].cells[j]
        set_cell_bg(cell, bg)
        align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(cell, val, size=8.5, align=align)

p_kol = doc.add_paragraph()
p_kol.paragraph_format.space_before = Pt(2)
rk = p_kol.add_run('📌 洞察：')
rk.bold = True; rk.font.size = Pt(8.5); rk.font.color.rgb = DARK_GREEN
rk2 = p_kol.add_run(
    ' 萬事如翊以最低成本（NT$3,000）達到最高觀看率（10%），CP值最佳。'
    '100天便當計畫費用最高（NT$75,000），但觀看率僅 0.3%，建議重新評估合作效益。'
    ' 免費合作 KOL（小廚娘阿貝、伊娃等）觀看率合計表現不亞於部分付費合作，未來可擴大免費換購模式。'
)
rk2.font.size = Pt(8.5)

# ══════════════════════════════════════
#  分頁
# ══════════════════════════════════════
doc.add_page_break()

# ══════════════════════════════════════
#  三、Meta 廣告投放分析
# ══════════════════════════════════════
add_section_title(doc, '三、Meta 廣告投放分析')

# Campaign A
p_ca = doc.add_paragraph()
p_ca.paragraph_format.space_before = Pt(3)
r_ca = p_ca.add_run('▌ 活動 A：全聯禮券抽獎活動－圖文廣告  （2026/06/06–07/05）')
r_ca.bold = True; r_ca.font.size = Pt(9.5); r_ca.font.color.rgb = MID_GREEN

ca_table = doc.add_table(rows=2, cols=6)
ca_table.style = 'Table Grid'
ca_headers = ['總曝光', '觸及人數', 'CTR', '每次成果成本', '總花費', 'Feed CTR']
ca_vals    = ['109,002', '56,770', '5.18%', 'NT$ 2', 'NT$ 6,238', '7.93%']
ca_col_w   = [Cm(2.3), Cm(2.3), Cm(1.5), Cm(2.3), Cm(2.0), Cm(1.9)]
for j, h in enumerate(ca_headers):
    cell = ca_table.rows[0].cells[j]
    cell.width = ca_col_w[j]
    set_cell_bg(cell, MID_GREEN)
    cell_text(cell, h, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for j, v in enumerate(ca_vals):
    cell = ca_table.rows[1].cells[j]
    set_cell_bg(cell, LIGHT_GREEN)
    cell_text(cell, v, bold=True, size=9, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)

p_ca2 = doc.add_paragraph()
p_ca2.paragraph_format.space_before = Pt(2)
p_ca2.paragraph_format.space_after  = Pt(1)
rc2 = p_ca2.add_run('版位：')
rc2.bold = True; rc2.font.size = Pt(8.5)
p_ca2.add_run(' FB 佔 99.2% 預算；Feed CTR 7.93% 遠優於 Reels（3.59%）。').font.size = Pt(8.5)
p_ca3 = doc.add_paragraph()
p_ca3.paragraph_format.space_before = Pt(0)
p_ca3.paragraph_format.space_after  = Pt(3)
rc3 = p_ca3.add_run('黃金時段：')
rc3.bold = True; rc3.font.size = Pt(8.5)
p_ca3.add_run(' 18:00–21:59 貢獻約 30% 花費，台北市（CTR 6.01%）為最佳地區。').font.size = Pt(8.5)

# Campaign B
p_cb = doc.add_paragraph()
p_cb.paragraph_format.space_before = Pt(4)
r_cb = p_cb.add_run('▌ 活動 B：全聯抽獎影片廣告  （2026/07/01 啟動，預計至 07/14）')
r_cb.bold = True; r_cb.font.size = Pt(9.5); r_cb.font.color.rgb = MID_GREEN

cb_table = doc.add_table(rows=2, cols=5)
cb_table.style = 'Table Grid'
cb_headers = ['曝光次數', 'ThruPlays', '貼文互動', 'CTR',  '每次成果成本']
cb_vals    = ['30,185', '5,247', '16,290', '8.14%', 'NT$ 1']
cb_col_w   = [Cm(2.5), Cm(2.2), Cm(2.2), Cm(1.8), Cm(2.5)]
for j, h in enumerate(cb_headers):
    cell = cb_table.rows[0].cells[j]
    cell.width = cb_col_w[j]
    set_cell_bg(cell, MID_GREEN)
    cell_text(cell, h, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for j, v in enumerate(cb_vals):
    cell = cb_table.rows[1].cells[j]
    set_cell_bg(cell, LIGHT_GREEN)
    cell_text(cell, v, bold=True, size=9, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)

p_cb2 = doc.add_paragraph()
p_cb2.paragraph_format.space_before = Pt(2)
p_cb2.paragraph_format.space_after  = Pt(1)
rcb2 = p_cb2.add_run('版位：')
rcb2.bold = True; rcb2.font.size = Pt(8.5)
p_cb2.add_run(' FB Reels 消耗 72.9% 預算；Feed CTR 高達 13.36%，建議提高 Feed 預算配比。').font.size = Pt(8.5)
p_cb3 = doc.add_paragraph()
p_cb3.paragraph_format.space_before = Pt(0)
p_cb3.paragraph_format.space_after  = Pt(4)
rcb3 = p_cb3.add_run('黃金時段：')
rcb3.bold = True; rcb3.font.size = Pt(8.5)
p_cb3.add_run(' 12:00–13:59 CTR 達 9.04%（午間最活躍）；單日預算 NT$800 未完全消耗，可評估調升競價。').font.size = Pt(8.5)

# ══════════════════════════════════════
#  四、綜合洞察與行動建議
# ══════════════════════════════════════
add_section_title(doc, '四、綜合洞察與行動建議')

recommendations = [
    ('🏆 主力產品強化',
     'YiBiYaYa 系列（果汁+益生菌）是流量與營收雙重主力，建議 7 月持續以此為廣告主素材，搭配抽獎活動延伸曝光。'),
    ('📱 廣告版位優化',
     'Feed 版位 CTR 持續優於 Reels（圖文廣告 7.93% vs 3.59%；影片廣告 13.36% vs Reels）。建議將 Feed 預算比例從目前 23–24% 提升至 40–50%。'),
    ('⏰ 投放時段聚焦',
     '圖文廣告以晚間 18–22 時為主、影片廣告以午間 12–14 時最佳。建議依廣告類型設定不同的時段出價調整（Bid Adjustment）。'),
    ('🤝 KOL 策略調整',
     '萬事如翊（NT$3,000 / 觀看率 10%）CP值最佳，建議深化長期合作。NT$75,000 的百萬KOL效益未達預期，下次合作前需重新議定以觀看率為保底指標的合約條款。'),
    ('💡 素材多元化',
     '影片廣告目前僅使用單一素材，建議在剩餘 7/14 到期前新增 1–2 組影片，避免受眾疲勞導致 CPR 上升。'),
]

rec_table = doc.add_table(rows=len(recommendations), cols=2)
rec_table.style = 'Table Grid'
rec_col_w = [Cm(3.5), Cm(10.2)]
for i, (title, content) in enumerate(recommendations):
    bg = LIGHT_GREEN if i % 2 == 0 else WHITE
    c0 = rec_table.rows[i].cells[0]
    c1 = rec_table.rows[i].cells[1]
    c0.width = rec_col_w[0]; c1.width = rec_col_w[1]
    set_cell_bg(c0, bg); set_cell_bg(c1, bg)
    cell_text(c0, title, bold=True, size=8.5, color=MID_GREEN)
    cell_text(c1, content, size=8.5)

# ── 頁尾 ──
footer_p = doc.add_paragraph()
footer_p.paragraph_format.space_before = Pt(8)
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rf = footer_p.add_run('十全特好食品股份有限公司  ·  行銷部  ·  2026.07.06')
rf.font.size = Pt(7.5); rf.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

output_path = '/home/user/-/十全_行銷月報_202606.docx'
doc.save(output_path)
print(f'Done: {output_path}')
